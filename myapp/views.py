from django.shortcuts import redirect, render
from .models import Document
from .forms import DocumentForm
from django.http import JsonResponse
from django.http import HttpResponse
from docx import Document as DOCX
from django.core import serializers
from dict2xml import dict2xml
import json



#Upload the Document file and save a copy to the database model Table

def uploadDocx(request):
    message = ''
    # Handle file upload
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            newdoc = Document(docfile=request.FILES['docfile'])
            newdoc.docfile.name='Table1.docx'
            newdoc.save()
            print(newdoc.docfile.url)
            message = 'File Uploaded Succesfuly'


            documents = Document.objects.all()

            # Redirect to the document list after POST
            return render(request,'list.html',{'documents': documents,'form': form,'message':message})
        else:
            message = 'The form is not valid. Fix the following error:'
    else:
        form = DocumentForm()  # An empty, unbound form

    # Load documents for the list page
    documents = Document.objects.all()

    # Render list page with the documents and the form
    context = {'documents': documents, 'form': form, 'message': message}
    return render(request, 'list.html', context)






 #Convert The selected table to a json object   

def  JsonTable(request):
    #f = open('media/documents/Table1.docx', 'rb')
    #document = DOCX(f)

    #first we should get all the tables titles
    Headers=[]
    doc=DOCX('media/documents/Table1.docx')
    for paragraph in doc.paragraphs:
        if paragraph.style.name=='Heading 1':
            Headers.append(paragraph.text)

    print (Headers)

    tables = doc.tables
    
    table=doc.tables[0]

    data = []
    title={}
    title["title"]="Table1"


    data.append(title)

    #print(list(enumerate(table.rows)))

    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)


        # Establish the mapping based on the first row
        # headers; these will become the keys of our dictionary
        if i == 0:
            
            
            keys = tuple(text)
            
            continue

        # Construct a dictionary for this row, mapping
        # keys to values for this row
        row_data = dict(zip(keys, text))
        
        data.append(row_data)
        

    print(data)

    #Convert dictionary object to Json Object
    json_object = json.dumps(data, indent = 4)   

    #doc1 = Document.objects.filter(id=32)
    return render(request, 'demo.html', {"json":json_object})



#convert the selected table to a XML Object

def  XMLTable(request):

    doc=DOCX('media/documents/demo.docx')
    tables = doc.tables
    print(tables)
    table=doc.tables[0]
    data = []
    title={}
    title["title"]="Table1"

    print(title)
    data.append(title)

    print("table is")
    print(table.rows[0].cells[0].text)
    print(list(enumerate(table.rows)))

    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        print("the text is")
        print(text)

        # Establish the mapping based on the first row
        # headers; these will become the keys of our dictionary
        if i == 0:
            print(i)
            
            keys = tuple(text)
            print(keys)
            continue

        # Construct a dictionary for this row, mapping
        # keys to values for this row
        row_data = dict(zip(keys, text))
        print(text)
        data.append(row_data)
        print(data)

    print(data)

    #Convert the file to XML Object 
    xml = dict2xml(data) 
    print(xml)

    return render(request, 'xml.html', {"xml":xml})

    
